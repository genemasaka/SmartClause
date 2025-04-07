from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF
import io

def convert_to_docx(content: str) -> bytes:
    """Convert text content to a formatted DOCX document."""
    doc = Document()
    sections = content.split('\n\n')
    
    for section in sections:
        if not section.strip():
            continue
        
        if section.isupper() or section.startswith("NOTICE OF"):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(section)
            run.bold = True
            run.font.size = Pt(14)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif section.startswith("Signed") or "Signature" in section:
            paragraph = doc.add_paragraph()
            paragraph.add_run("\n" + "_" * 40 + "\n").bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            doc.add_paragraph(section)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

class LegalPDF(FPDF):
    """Custom PDF class for formatting legal documents."""
    def header(self):
        pass
    
    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")
    
    def add_title(self, text):
        self.set_font("Arial", "B", 14)
        self.cell(0, 10, text, ln=True, align="C")
        self.ln(5)
    
    def add_section(self, text):
        self.set_font("Arial", "", 11)
        self.multi_cell(0, 6, text)
        self.ln(3)
    
    def add_signature_line(self):
        self.ln(10)
        self.cell(70, 10, "_" * 30, ln=True)

def convert_to_pdf(content: str) -> bytes:
    """Convert text content to a formatted PDF document."""
    pdf = LegalPDF()
    pdf.add_page()
    sections = content.split('\n\n')
    
    for section in sections:
        if not section.strip():
            continue
        
        if section.isupper() or section.startswith("NOTICE OF"):
            pdf.add_title(section)
        elif section.startswith("Signed") or "Signature" in section:
            pdf.add_signature_line()
        else:
            pdf.add_section(section)
    
    return pdf.output(dest="S").encode("latin-1")
