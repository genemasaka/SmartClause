import streamlit as st
import logging
import json
import uuid
from pathlib import Path
from typing import Optional
import os
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from agent import setup_environment, DocumentGenerator
from mpesa_handler import MpesaHandler
from font_setup import setup_fonts
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import io
import re
import time
import html
import bleach
from payment_verification import PaymentVerification, init_payment_state, update_payment_status, handle_download_request, reset_payment_state
import base64

# Configure structured logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)
handler = logging.StreamHandler()
handler.setFormatter(logging.Formatter(
    '{"timestamp": "%(asctime)s", "name": "%(name)s", "level": "%(levelname)s", "message": %(message)s}'
))
logger.handlers = [handler]

def init_session_state():
    """Initialize session state variables"""
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "generator" not in st.session_state:
        try:
            template_manager, model = setup_environment()
            st.session_state.generator = DocumentGenerator(template_manager, model)
            logger.info('{"event": "generator_initialized", "status": "success"}')
        except Exception as e:
            logger.error('{"event": "generator_init_failed", "error": "%s"}', str(e), exc_info=True)
            st.error("Unable to initialize the application. Please try again later. If this issue persists, please report it using the Feedback tab in the Help Guide.")
    if "mpesa" not in st.session_state:
        st.session_state.mpesa = MpesaHandler()
        logger.info('{"event": "mpesa_handler_initialized", "status": "success"}')
    if "document_generated_successfully" not in st.session_state:
        st.session_state.document_generated_successfully = False
    if "show_welcome" not in st.session_state:
        st.session_state.show_welcome = True
    if "current_document" not in st.session_state:
        st.session_state.current_document = None
    if "generation_in_progress" not in st.session_state:
        st.session_state.generation_in_progress = False
    if "document_price" not in st.session_state:
        st.session_state.document_price = 0
    if "show_payment" not in st.session_state:
        st.session_state.show_payment = False
    if "show_download" not in st.session_state:
        st.session_state.show_download = False
    if "force_sidebar_open" not in st.session_state:
        st.session_state.force_sidebar_open = False
    if "session_id" not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
    if not setup_fonts():
        st.warning("Some fonts are missing. Documents may not display correctly.")
    init_payment_state()

def validate_input(input_value, input_type):
    """
    Enhanced validation with additional security checks
    """
    if input_type == 'phone':
        # Existing phone validation
        return input_value.startswith("254") and len(input_value) == 12 and input_value.isdigit()
    
    elif input_type == 'email':
        # Enhanced email validation with additional security checks
        if not input_value:
            return False
        
        # Check for injection attempts
        dangerous_chars = ['\n', '\r', '\t', '<', '>', '"', "'"]
        if any(char in input_value for char in dangerous_chars):
            logger.warning("Potentially malicious email input blocked")
            return False
        
        # Basic email regex
        import re
        pattern = r'^[\w\.-]+@[\w\.-]+\.\w+$'
        return bool(re.match(pattern, input_value))
    
    elif input_type == 'prompt':
        if not input_value or len(input_value) > 1000 or len(input_value.strip()) < 5:
            return False
        
        # Additional check for excessive special characters (potential obfuscation)
        special_char_ratio = sum(1 for c in input_value if not c.isalnum() and not c.isspace()) / len(input_value)
        if special_char_ratio > 0.3:  # More than 30% special characters
            logger.warning("Input with excessive special characters blocked")
            return False
        
        return True
    
    elif input_type == 'feedback':
        if len(input_value) > 2000:
            return False
        return True
    
    return False

def sanitize_text(text):
    """
    Enhanced sanitize text input to remove potential malicious content
    Fixes issues with JavaScript URLs and whitespace handling
    """
    if not text:
        return text
    
    # First, remove javascript: URLs and other dangerous protocols
    dangerous_protocols = [
        r'javascript\s*:',
        r'vbscript\s*:',
        r'data\s*:',
        r'about\s*:',
    ]
    
    cleaned_text = text
    for protocol in dangerous_protocols:
        cleaned_text = re.sub(protocol, '', cleaned_text, flags=re.IGNORECASE)
    
    # Remove HTML tags using bleach
    cleaned_text = bleach.clean(cleaned_text, tags=[], strip=True)
    
    # Clean up extra whitespace that might be left from nested tags
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text.strip())
    
    # Additional safety: Remove any remaining event handlers
    dangerous_attributes = [
        r'on\w+\s*=',  # onclick, onload, onerror, etc.
        r'href\s*=\s*["\']?\s*javascript:',
        r'src\s*=\s*["\']?\s*javascript:',
    ]
    
    for attr in dangerous_attributes:
        cleaned_text = re.sub(attr, '', cleaned_text, flags=re.IGNORECASE)
    
    return cleaned_text

class UnicodeAwarePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_font('DejaVu', '', 'DejaVuSansCondensed.ttf', uni=True)
        self.add_font('DejaVu', 'B', 'DejaVuSansCondensed-Bold.ttf', uni=True)
        self.set_font('DejaVu', size=11)
        self.set_auto_page_break(auto=True, margin=15)
    
    def multi_cell(self, w, h, txt, border=0, align='L', fill=False):
        if w == 0:
            w = self.w - self.l_margin - self.r_margin
        wmax = (w - 2 * self.c_margin) * 1000 / self.font_size
        text = txt.replace('\r', '')
        super().multi_cell(w, h, text, border, align, fill)

class KenyanLegalDocument:
    def clean_text_for_pdf(text: str) -> str:
        """Clean and prepare text for PDF conversion by replacing problematic Unicode characters."""
        replacements = {
            '\u2013': '-',  # en dash
            '\u2014': '--', # em dash
            '\u2018': "'",  # left single quote
            '\u2019': "'",  # right single quote
            '\u201C': '"',  # left double quote
            '\u201D': '"',  # right double quote
            '\u2026': '...', # ellipsis
            '\u2022': '*',  # bullet point
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        cleaned_text = ''
        for char in text:
            try:
                char.encode('latin-1')
                cleaned_text += char
            except UnicodeEncodeError:
                cleaned_text += ' '
                
        return cleaned_text

class LegalDocumentPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_page()
        self.set_margins(25, 25, 25)
        self.set_auto_page_break(auto=True, margin=25)
    
    def chapter_title(self, text):
        """Add a chapter title with specific formatting"""
        self.set_font('Helvetica', 'BU', 13)
        self.cell(0, 10, text, ln=True, align='C')
        self.ln(10)
    
    def body_text(self, text):
        """Add body text with standard formatting"""
        self.set_font('Helvetica', '', 11)
        self.multi_cell(0, 10, text)
        self.ln(5)

def clean_markdown(text: str) -> tuple[str, bool]:
    """Remove markdown symbols except square brackets, escape HTML, and return the clean text and whether it was bold"""
    is_bold = False
    cleaned_text = text
    
    if '**' in text:
        cleaned_text = text.replace('**', '')
        is_bold = True
    
    cleaned_text = re.sub(r'\\|\*', '', cleaned_text)
    cleaned_text = cleaned_text.replace('{.underline}', '')
    
    # Fix common AI-generated spelling errors in legal documents
    # These corrections should happen BEFORE any other processing
    spelling_corrections = {
        "THEREFOREE": "THEREFORE",
        "WHEREOFF": "WHEREOF",
        "NOW, THEREFOREE": "NOW, THEREFORE",
        "IN WITNESS WHEREOFF": "IN WITNESS WHEREOF"
    }
    
    for incorrect, correct in spelling_corrections.items():
        cleaned_text = cleaned_text.replace(incorrect, correct)
    
    escaped_text = html.escape(cleaned_text.strip())
    return escaped_text, is_bold

def identify_document_type(content: str) -> str:
    """Identify if the document is an affidavit or contract"""
    if "OATHS AND STATUTORY DECLARATIONS ACT" in content:
        return "affidavit"
    elif "CONTRACT" in content.upper() or "AGREEMENT" in content.upper():
        return "contract"
    return "other"

def format_affidavit_docx(doc, content: str):
    """Apply specific formatting for affidavit documents in DOCX"""
    title_sections = [
        "REPUBLIC OF KENYA",
        "IN THE MATTER OF THE OATHS AND STATUTORY DECLARATIONS ACT",
        "(CAP 15 OF THE LAWS OF KENYA)",
        "AFFIDAVIT"
    ]
    
    for title in title_sections:
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.underline = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    paragraphs = content.split('\n\n')
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
            
        if any(title in para_text for title in title_sections):
            continue
        
        para = doc.add_paragraph()
        
        if "THAT" in para_text:
            match = re.match(r'(\d+\.\s+)?(\*\*)?THAT(\*\*)?(.+)', para_text)
            if match:
                number = match.group(1) or ''
                remainder = match.group(4)
                
                if number:
                    para.add_run(number)
                
                that_run = para.add_run("THAT ")
                that_run.bold = True
                that_run.font.underline = True
                
                clean_remainder, is_bold = clean_markdown(remainder)
                remainder_run = para.add_run(clean_remainder)
                if is_bold:
                    remainder_run.bold = True
        
        elif "SWORN" in para_text:
            lines = para_text.split('\n')
            for line in lines:
                clean_line, is_bold = clean_markdown(line)
                if clean_line:
                    sworn_para = doc.add_paragraph()
                    run = sworn_para.add_run(clean_line)
                    run.bold = True
                    sworn_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        else:
            clean_text, is_bold = clean_markdown(para_text)
            run = para.add_run(clean_text)
            if is_bold:
                run.bold = True

def format_contract_docx(doc, content: str):
    """Apply specific formatting for contract documents in DOCX"""
    content = content.replace("THEREFOREE", "THEREFORE")
    content = content.replace("WHEREOFF", "WHEREOF")
    paragraphs = content.split('\n\n')
    title_added = False
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
            
        para = doc.add_paragraph()
        
        if not title_added and ("CONTRACT" in para_text.upper() or "AGREEMENT" in para_text.upper()):
            title_para = doc.add_paragraph()
            title_run = title_para.add_run("CONTRACT AGREEMENT")
            title_run.bold = True
            title_run.font.underline = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_added = True
            continue
        
        if re.match(r'^ARTICLE\s+\d+:', para_text, re.IGNORECASE):
            parts = para_text.split(':', 1)
            if len(parts) == 2:
                title, content = parts
                run = para.add_run(title + ':')
                run.bold = True
                para.add_run(content)
            else:
                run = para.add_run(para_text)
                run.bold = True
        elif "WHEREAS:" in para_text:
            para.add_run("WHEREAS:").bold = True
            para.add_run(para_text[8:])
        elif "NOW, THEREFORE" in para_text:
            para.add_run("NOW, THEREFORE").bold = True
            para.add_run(para_text[13:])
        elif "IN WITNESS WHEREOF" in para_text:
            para.add_run("IN WITNESS WHEREOF").bold = True
            para.add_run(para_text[17:])
        else:
            clean_text, is_bold = clean_markdown(para_text)
            run = para.add_run(clean_text)
            if is_bold:
                run.bold = True

def convert_to_docx(content: str) -> bytes:
    """Convert text content to DOCX format with appropriate formatting"""
    session_id = st.session_state.get("session_id", "unknown")
    try:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        doc_type = identify_document_type(content)
        
        if doc_type == "affidavit":
            format_affidavit_docx(doc, content)
        elif doc_type == "contract":
            format_contract_docx(doc, content)
        else:
            for para in content.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        buffer = io.BytesIO()
        doc.save(buffer)
        logger.info('{"event": "docx_conversion", "session_id": "%s", "doc_type": "%s", "status": "success"}', session_id, doc_type)
        return buffer.getvalue()
    except Exception as e:
        logger.error('{"event": "docx_conversion_failed", "session_id": "%s", "error": "%s"}', session_id, str(e), exc_info=True)
        st.error("Failed to generate DOCX document. Please try again later. If this issue persists, please report it using the Feedback tab in the Help Guide.")
        return None

def format_affidavit_pdf(pdf: FPDF, content: str):
    """Apply specific formatting for affidavit documents in PDF"""
    title_sections = [
        "REPUBLIC OF KENYA",
        "IN THE MATTER OF THE OATHS AND STATUTORY DECLARATIONS ACT",
        "(CAP 15 OF THE LAWS OF KENYA)",
        "AFFIDAVIT"
    ]
    
    for title in title_sections:
        pdf.set_font('Helvetica', 'BU', 13)
        pdf.cell(0, 10, title, ln=True, align='C')
    
    pdf.ln(10)
    
    paragraphs = content.split('\n\n')
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
            
        if any(title in para_text for title in title_sections):
            continue
        
        if "THAT" in para_text:
            match = re.match(r'(\d+\.\s+)?(\*\*)?THAT(\*\*)?(.+)', para_text)
            if match:
                number = match.group(1) or ''
                remainder = match.group(4)
                
                pdf.set_font('Helvetica', '', 12)
                if number:
                    pdf.cell(10, number)
                
                pdf.set_font('Helvetica', 'BU', 12)
                pdf.cell(10, "THAT ")
                
                clean_remainder, is_bold = clean_markdown(remainder)
                pdf.set_font('Helvetica', 'B' if is_bold else '', 12)
                pdf.multi_cell(0, 10, clean_remainder)
        
        elif "SWORN" in para_text:
            lines = para_text.split('\n')
            for line in lines:
                clean_line, _ = clean_markdown(line)
                if clean_line:
                    pdf.set_font('Helvetica', 'B', 12)
                    pdf.cell(0, 10, clean_line, ln=True)
        
        else:
            clean_text, is_bold = clean_markdown(para_text)
            pdf.set_font('Helvetica', 'B' if is_bold else '', 12)
            pdf.multi_cell(0, 10, clean_text)
        
        pdf.ln(5)

def format_contract_pdf(pdf: FPDF, content: str):
    """Apply specific formatting for contract documents in PDF"""
    content = content.replace("THEREFOREE", "THEREFORE")
    content = content.replace("WHEREOFF", "WHEREOF")
    paragraphs = content.split('\n\n')
    title_added = False
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
            
        if not title_added and ("CONTRACT" in para_text.upper() or "AGREEMENT" in para_text.upper()):
            pdf.set_font('Helvetica', 'BU', 14)
            pdf.cell(0, 10, "CONTRACT AGREEMENT", ln=True, align='C')
            pdf.ln(10)
            title_added = True
            continue
        
        if re.match(r'^ARTICLE\s+\d+:', para_text, re.IGNORECASE):
            parts = para_text.split(':', 1)
            if len(parts) == 2:
                title, content = parts
                pdf.set_font('Helvetica', 'B', 12)
                pdf.cell(10, title + ':')
                pdf.set_font('Helvetica', '', 12)
                pdf.multi_cell(0, 10, content)
            else:
                pdf.set_font('Helvetica', 'B', 12)
                pdf.multi_cell(0, 10, para_text)
        elif "WHEREAS:" in para_text:
            pdf.set_font('Helvetica', 'B', 12)
            pdf.cell(10, "WHEREAS:")
            pdf.set_font('Helvetica', '', 12)
            pdf.multi_cell(0, 10, para_text[8:])
        elif "NOW, THEREFORE" in para_text:
            pdf.set_font('Helvetica', 'B', 12)
            pdf.cell(10, "NOW, THEREFORE")
            pdf.set_font('Helvetica', '', 12)
            pdf.multi_cell(0, 10, para_text[13:])
        elif "IN WITNESS WHEREOF" in para_text:
            pdf.set_font('Helvetica', 'B', 12)
            pdf.cell(10, "IN WITNESS WHEREOF")
            pdf.set_font('Helvetica', '', 12)
            pdf.multi_cell(0, 10, para_text[17:])
        else:
            clean_text, is_bold = clean_markdown(para_text)
            pdf.set_font('Helvetica', 'B' if is_bold else '', 12)
            pdf.multi_cell(0, 10, clean_text, align='L')
        
        pdf.ln(5)

def convert_to_pdf(content: str) -> bytes:
    """Convert document content to PDF with appropriate formatting"""
    session_id = st.session_state.get("session_id", "unknown")
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_margins(25, 25, 25)
        pdf.set_auto_page_break(auto=True, margin=25)
        
        doc_type = identify_document_type(content)
        
        if doc_type == "affidavit":
            format_affidavit_pdf(pdf, content)
        elif doc_type == "contract":
            format_contract_pdf(pdf, content)
        else:
            pdf.set_font('Helvetica', '', 12)
            for para in content.split('\n\n'):
                if para.strip():
                    pdf.multi_cell(0, 10, para.strip())
                    pdf.ln(5)
        
        pdf_data = pdf.output(dest='S').encode('latin-1')
        logger.info('{"event": "pdf_conversion", "session_id": "%s", "doc_type": "%s", "status": "success"}', session_id, doc_type)
        return pdf_data
    except Exception as e:
        logger.error('{"event": "pdf_conversion_failed", "session_id": "%s", "error": "%s"}', session_id, str(e), exc_info=True)
        st.error("Failed to generate PDF document. Please try again later. If this issue persists, please report it using the Feedback tab in the Help Guide.")
        return None

def show_download_buttons():
    """Display persistent download buttons for the document"""
    session_id = st.session_state.get("session_id", "unknown")
    if st.session_state.show_download and st.session_state.current_document:
        payment_verifier = PaymentVerification(st.session_state.mpesa)
        
        try:
            if not st.session_state.payment_verified:
                st.warning("Please open the side bar and complete payment to download file.")
            else:
                st.success("✅ Payment verified")
            
            st.markdown("### Download Options")
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("Download as DOCX", key="download_docx_button"):
                    if handle_download_request(payment_verifier):
                        docx_data = convert_to_docx(st.session_state.current_document)
                        if docx_data:
                            st.download_button(
                                label="Click to download DOCX",
                                data=docx_data,
                                file_name="document.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="docx_download"
                            )
                        else:
                            logger.error('{"event": "docx_download_failed", "session_id": "%s", "reason": "conversion_failed"}', session_id)
            
            with col2:
                if st.button("Download as PDF", key="download_pdf_button"):
                    if handle_download_request(payment_verifier):
                        pdf_data = convert_to_pdf(st.session_state.current_document)
                        if pdf_data:
                            st.download_button(
                                label="Click to download PDF",
                                data=pdf_data,
                                file_name="document.pdf",
                                mime="application/pdf",
                                key="pdf_download"
                            )
                        else:
                            logger.error('{"event": "pdf_download_failed", "session_id": "%s", "reason": "conversion_failed"}', session_id)
                
        except Exception as e:
            logger.error('{"event": "download_preparation_failed", "session_id": "%s", "error": "%s"}', session_id, str(e), exc_info=True)
            st.error("An error occurred while preparing the download. Please try again later. If this issue persists, please report it using the Feedback tab in the Help Guide.")

def generate_document(generator: DocumentGenerator, prompt: str) -> tuple[Optional[str], bool]:
    session_id = st.session_state.get("session_id", "unknown")
    try:
        response = generator.generate_document(prompt)
        if response:
            st.session_state.current_document = response
            st.session_state.document_price = DocumentPricing.get_document_price(response)
            reset_payment_state()
            logger.info('{"event": "document_generated", "session_id": "%s", "doc_type": "%s", "status": "success"}', 
                        session_id, identify_document_type(response))
            return response, True
        else:
            logger.error('{"event": "document_generation_failed", "session_id": "%s", "reason": "no_response"}', session_id)
            return "Unable to generate the document. Please try again with a different prompt. If this issue persists, please report it using the Feedback tab in the Help Guide.", False
    except Exception as e:
        logger.error('{"event": "document_generation_error", "session_id": "%s", "error": "%s"}', session_id, str(e), exc_info=True)
        return "An error occurred while generating the document. Please try again later. If this issue persists, please report it using the Feedback tab in the Help Guide.", False

def validate_phone_number(phone: str) -> bool:
    """Validate the phone number format"""
    session_id = st.session_state.get("session_id", "unknown")
    if phone.startswith("254") and len(phone) == 12 and phone.isdigit():
        return True
    logger.warning('{"event": "invalid_phone_number", "session_id": "%s", "phone_length": %d}', session_id, len(phone))
    return False

class DocumentPricing:
    """Handles pricing logic for different document types"""
    PRICES = {
        "affidavit": 3000,
        "contract": 5000,
        "other": 5000
    }
    
    @staticmethod
    def get_document_price(content: str) -> int:
        if "OATHS AND STATUTORY DECLARATIONS ACT" in content:
            return DocumentPricing.PRICES["affidavit"]
        elif "CONTRACT" in content.upper() or "AGREEMENT" in content.upper():
            return DocumentPricing.PRICES["contract"]
        return DocumentPricing.PRICES["other"]

def show_welcome_modal():
    """Shows a welcome modal with updated quick start guide"""
    if st.session_state.show_welcome:
        with st.container():
            st.markdown(""" 
            # Welcome!
            
            Let's help you get started with generating your legal documents.
            """)
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown(""" 
                ### Quick Start
                1. Type your request naturally
                2. Wait for document generation
                3. Review the document
                4. Make payment via M-PESA
                5. Download your document
                
                ### Example Prompts
                - "Draft a contract for web development services"
                - "Create an affidavit for change of name"
                - "Generate a stock transfer agreement"
                """)
                
            with col2:
                st.markdown(""" 
                ### Pro Tips
                - Be specific about parties involved
                - Include key terms and conditions
                - Mention important dates
                - Include any special requirements
                
                ### ⚠️ Important Notes
                - All documents should be reviewed by a legal professional
                - Sensitive information is encrypted
                - Documents are compliant with Kenyan law
                """)

            if not st.session_state.get("generation_in_progress", False):
                if st.button("Got it!", key="welcome_close"):
                    st.session_state.show_welcome = False
                    if not st.session_state.get("generation_in_progress", False):
                        st.rerun()
            else:
                st.button("Got it!", key="welcome_close", disabled=True)

def send_feedback_email(feedback: str, user_email: str = ""):
    """Send feedback email to smartclause6@gmail.com"""
    session_id = st.session_state.get("session_id", "unknown")
    email_address = os.environ.get('EMAIL_ADDRESS')
    email_password = os.environ.get('EMAIL_PASSWORD')
    
    if not email_address or not email_password:
        logger.error('{"event": "email_config_missing", "session_id": "%s"}', session_id)
        raise ValueError("Email configuration is missing. Please contact support.")
    
    msg = MIMEMultipart()
    msg['From'] = email_address
    msg['To'] = "smartclause6@gmail.com"
    msg['Subject'] = "User Feedback"
    
    body = f"Feedback:\n{feedback}\n\nFrom: {user_email if user_email else 'Anonymous'}"
    msg.attach(MIMEText(body, 'plain'))
    
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(email_address, email_password)
        server.send_message(msg)
        server.quit()
        logger.info('{"event": "feedback_email_sent", "session_id": "%s", "status": "success"}', session_id)
    except Exception as e:
        logger.error('{"event": "feedback_email_failed", "session_id": "%s", "error": "%s"}', session_id, str(e), exc_info=True)
        raise Exception("Failed to send feedback. Please try again later. If this issue persists, please report it using the Feedback tab in the Help Guide.")

def show_help_section():
    """Shows the comprehensive help section with feedback tab"""
    with st.expander("Help Guide", expanded=False):
        tabs = st.tabs(["Document Types", "How to Use", "Best Practices", "FAQ", "Feedback"])
        
        with tabs[0]:
            st.markdown("""
            ### Available Document Types
            
            1. **Contract** - For business agreements and services
            2. **Affidavit** - Sworn statements for legal proceedings
            
            Each document type follows specific Kenyan legal requirements and formatting.
            """)
            
        with tabs[1]:
            st.markdown("""
            ### Step-by-Step Guide
            
            1. **Formulate Your Request**
               - Simply describe what you need
               - Example: "Create a contract for a house sale"
               - Include all relevant details
               - Be specific about requirements
            
            2. **Review Generated Document**
               - Check all details are correct
               - Verify names and dates
               - Ensure terms match your needs
            
            3. **Make Payment**
               - Enter your M-PESA number
               - Confirm payment prompt on your phone
               - Complete the transaction
            
            4. **Download and Use**
               - Download the final document
               - Review with legal counsel
               - Make any necessary modifications
            """)
            
        with tabs[2]:
            st.markdown("""
            ### Best Practices for Document Generation
            
            1. **Be Specific in Your Request**
               - Include full names of all parties
               - Specify dates and deadlines
               - Detail all terms and conditions
               - Mention specific requirements
            
            2. **Include Essential Details**
               - Business/Company names
               - Registration numbers
               - Physical addresses
               - Contact information
               - Monetary values
               - Time periods
            
            3. **Consider Legal Requirements**
               - Jurisdiction specifications
               - Regulatory compliance needs
               - Required certifications
               - Witness requirements
            
            4. **Document Review**
               - Always read thoroughly
               - Check for accuracy
               - Verify all parties' details
               - Confirm terms and conditions
            """)
            
        with tabs[3]:
            st.markdown("""
            ### Frequently Asked Questions
            
            **Q: Are these documents legally binding?**  
            A: While generated documents follow legal formats, they should be reviewed by a qualified legal professional before use.
            
            **Q: How secure is my information?**  
            A: Your information is held temporarily in memory during your active session and is not stored persistently by the app after the session ends. 
               Data transmitted to external services (e.g., M-PESA for payments and the AI model for document generation) is sent over secure connections, 
               but storage by these services is subject to their privacy policies. We recommend reviewing M-PESA and Google’s (Gemini AI) privacy terms for details on how they handle your data.
            
            **Q: Can I modify the generated document?**  
            A: Yes, you can request modifications by providing specific changes needed in your prompt.
            
            **Q: What if I need a document type not listed?**  
            A: Contact support for custom document types or use the closest matching document type.
            
            **Q: How long does generation take?**  
            A: Document generation typically takes 30-60 seconds depending on complexity.
            
            **Q: What payment methods are accepted?**  
            A: Currently, we accept M-PESA payments only.
                        
            📱 Need Help?
                Email: smartclause6@gmail.com
                Call: +254 XXX XXX XXX
            """)
            
        with tabs[4]:
            st.subheader("Send us your feedback")
            with st.form("feedback_form"):
                feedback = st.text_area("Your feedback:", height=150)
                user_email = st.text_input("Your email (optional):")
                submitted = st.form_submit_button("Send Feedback")
                
                if submitted:
                    if feedback:
                        if user_email and not validate_input(user_email, 'email'):
                            st.error("Invalid email format. If this issue persists, please report it using the Feedback tab in the Help Guide.")
                        elif not validate_input(feedback, 'feedback'):
                            st.error("Feedback must be less than 2000 characters. If this issue persists, please report it using the Feedback tab in the Help Guide.")
                        else:
                            sanitized_feedback = sanitize_text(feedback)
                            try:
                                send_feedback_email(sanitized_feedback, user_email)
                                st.success("Thank you for your feedback!")
                            except Exception as e:
                                st.error("Failed to send feedback. Please try again later. If this issue persists, please report it using the Feedback tab in the Help Guide.")
                    else:
                        st.warning("Please enter your feedback before submitting.")

def enhance_sidebar():
    """Enhanced sidebar with additional information"""
    with st.sidebar:
        st.header("Document Quality Tips")
        st.info("""
        To get the best results:
        - Use clear, specific language
        - Include all relevant parties
        - Include all important details(eg. ID numbers, Postal addresses, vehicle registration numbers, title deed numbers, etc)
        - Specify important dates
        - Detail key terms
        """)

        st.header("Controls")
        if not st.session_state.get("generation_in_progress", False):
            if st.button("Clear Chat History", key="clear_chat_button"):
                st.session_state.messages = []
                st.session_state.show_payment = False
                st.session_state.document_generated_successfully = False
                st.session_state.show_download = False
                st.session_state.current_document = None
                st.session_state.force_sidebar_open = False
                logger.info('{"event": "chat_history_cleared", "session_id": "%s"}', st.session_state.get("session_id", "unknown"))
                st.rerun()
        else:
            st.button("Clear Chat History", key="clear_chat_button", disabled=True)

        if st.session_state.get("show_payment", False) and st.session_state.get("document_generated_successfully", False):
            st.markdown('<div class="payment-card">', unsafe_allow_html=True)
            with st.form("sidebar_payment_form"):
                st.subheader("💳 Pay using M-PESA STK")
                st.write(f"Amount to pay: KES {st.session_state.document_price:,}")
                phone = st.text_input("Phone number (254XXXXXXXXX):")
                submitted = st.form_submit_button("Pay Now")
                
                if submitted and phone:
                    session_id = st.session_state.get("session_id", "unknown")
                    if validate_phone_number(phone):
                        amount = st.session_state.document_price
                        account_ref = "Smart Clause"
                        desc = "Payment for legal document generation"
                        try:
                            response = st.session_state.mpesa.initiate_stk_push(phone, amount, desc)
                            if 'ResponseCode' in response and response['ResponseCode'] == '0':
                                update_payment_status(response['CheckoutRequestID'], amount)
                                st.success("✓ Check your phone for payment prompt")
                                logger.info('{"event": "payment_initiated", "session_id": "%s", "checkout_request_id": "%s", "amount": %d}', 
                                            session_id, response['CheckoutRequestID'], amount)
                            else:
                                st.error("✗ Payment initiation failed. Please try again. If this issue persists, please report it using the Feedback tab in the Help Guide.")
                                logger.error('{"event": "payment_initiation_failed", "session_id": "%s", "error": "%s"}', 
                                             session_id, response.get('errorMessage', 'Unknown error'))
                        except Exception as e:
                            st.error("✗ An error occurred during payment. Please try again later. If this issue persists, please report it using the Feedback tab in the Help Guide.")
                            logger.error('{"event": "payment_initiation_error", "session_id": "%s", "error": "%s"}', session_id, str(e), exc_info=True)
                    else:
                        st.error("✗ Invalid phone number format. Use 254XXXXXXXXX. If this issue persists, please report it using the Feedback tab in the Help Guide.")
            st.markdown('</div>', unsafe_allow_html=True)

def format_document_html(content: str) -> str:
    """Convert document content to formatted HTML for display in the web interface with theme awareness"""
    session_id = st.session_state.get("session_id", "unknown")
    logger.info('{"event": "format_document_html", "session_id": "%s", "content_length": %d}', session_id, len(content))

    doc_type = identify_document_type(content)
    
    html = f"""
    <div class="legal-document" style="
        font-family: 'Times New Roman', serif; 
        line-height: 1.5; 
        padding: 25px; 
        max-width: 800px; 
        margin: 0 auto; 
        background-color: var(--document-bg-color, #ffffff); 
        color: var(--document-text-color, #000000); 
        border: 1px solid rgba(128, 128, 128, 0.2); 
        overflow: visible !important; 
        min-height: 100px;
        display: block;
    ">
    <style>
        .legal-document * {{
            max-height: none !important;
            overflow: visible !important;
        }}
    </style>
    """
    
    if doc_type == "affidavit":
        title_sections = [
            "REPUBLIC OF KENYA",
            "IN THE MATTER OF THE OATHS AND STATUTORY DECLARATIONS ACT",
            "(CAP 15 OF THE LAWS OF KENYA)",
            "AFFIDAVIT",
        ]
        
        for title in title_sections:
            html += f'<h6 class="doc-title" style="text-align: center; font-weight: bold; text-decoration: underline; color: var(--document-text-color, #000000); opacity: 1;">{title}</h6>'
        
        html += '<div style="margin-top: 14px; "></div>'
        
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
                
            if any(title in para_text for title in title_sections):
                continue
            
            if "THAT" in para_text:
                match = re.match(r'(\d+\.\s+)?(\*\*)?THAT(\*\*)?(.+)', para_text)
                if match:
                    number = match.group(1) or ''
                    remainder = match.group(4)
                    
                    clean_remainder, is_bold = clean_markdown(remainder)
                    
                    html += '<p class="doc-paragraph">'
                    if number:
                        html += f'{number}'
                    
                    html += '<span style="font-weight: bold; text-decoration: underline;">THAT</span> '
                    
                    if is_bold:
                        html += f'<span style="font-weight: bold;">{clean_remainder}</span>'
                    else:
                        html += clean_remainder
                    
                    html += '</p>'
            
            elif "SWORN" in para_text:
                lines = para_text.split('\n')
                for line in lines:
                    clean_line, _ = clean_markdown(line)
                    if clean_line:
                        html += f'<p class="doc-paragraph" style="font-weight: bold;">{clean_line}</p>'
            
            else:
                clean_text, is_bold = clean_markdown(para_text)
                if is_bold:
                    html += f'<p class="doc-paragraph" style="font-weight: bold;">{clean_text}</p>'
                else:
                    html += f'<p class="doc-paragraph">{clean_text}</p>'
    
    elif doc_type == "contract":
        paragraphs = content.split('\n\n')
        title_added = False
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
                
            if not title_added and ("CONTRACT" in para_text.upper() or "AGREEMENT" in para_text.upper()):
                html += '<h6 class="doc-title" style="text-align: center; font-weight: bold; text-decoration: underline; color: var(--document-text-color, #000000); opacity: 1;">CONTRACT AGREEMENT</h6>'
                title_added = True
                continue
            
            if re.match(r'^ARTICLE\s+\d+:', para_text, re.IGNORECASE):
                parts = para_text.split(':', 1)
                if len(parts) == 2:
                    title, content = parts
                    html += f'<p class="doc-paragraph"><span style="font-weight: bold;">{title}:</span>{content}</p>'
                else:
                    html += f'<p class="doc-paragraph" style="font-weight: bold;">{para_text}</p>'
            elif "WHEREAS:" in para_text:
                html += f'<p class="doc-paragraph"><span style="font-weight: bold;">WHEREAS:</span>{para_text[8:]}</p>'
            elif "NOW, THEREFORE" in para_text:
                html += f'<p class="doc-paragraph"><span style="font-weight: bold;">NOW, THEREFORE</span>{para_text[13:]}</p>'
            elif "IN WITNESS WHEREOF" in para_text:
                html += f'<p class="doc-paragraph"><span style="font-weight: bold;">IN WITNESS WHEREOF</span>{para_text[17:]}</p>'
            else:
                clean_text, is_bold = clean_markdown(para_text)
                if is_bold:
                    html += f'<p class="doc-paragraph" style="font-weight: bold;">{clean_text}</p>'
                else:
                    html += f'<p class="doc-paragraph">{clean_text}</p>'
    
    else:
        for para in content.split('\n\n'):
            if para.strip():
                html += f'<p class="doc-paragraph">{para.strip()}</p>'
    
    html += '</div>'
    return html

def show_main_content():
    
    st.markdown('<p style="font-size: 16px; color: #888888;">Generate professional legal documents compliant with Kenyan law</p>', unsafe_allow_html=True)

    if st.session_state.show_welcome:
        show_welcome_modal()
    show_help_section()
    
    chat_container = st.container()
    
    with chat_container:
        for message in st.session_state.messages:
            if message["role"] == "user" or (message["role"] == "assistant" and not st.session_state.document_generated_successfully):
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])
    
    if st.session_state.document_generated_successfully and st.session_state.current_document:
        st.markdown('<div class="document-preview">', unsafe_allow_html=True)
        st.markdown(format_document_html(st.session_state.current_document), unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    
    if st.session_state.show_download and st.session_state.current_document:
        show_download_buttons()

    if prompt := st.chat_input("Describe your legal document needs, e.g., 'Draft a contract for...'"):
        session_id = st.session_state.get("session_id", "unknown")
        if st.session_state.show_welcome:
            st.session_state.show_welcome = False
            
        with st.chat_message("user"):
            st.markdown(prompt)

        st.session_state.messages.append({"role": "user", "content": prompt})

        with st.chat_message("assistant"):
            response_placeholder = st.empty()
            
            st.session_state.generation_in_progress = True
            
            try:
                with st.spinner("Generating your document..."):
                    if validate_input(prompt, 'prompt'):
                        sanitized_prompt = sanitize_text(prompt)
                        preview, success = generate_document(st.session_state.generator, sanitized_prompt)
                    else:
                        preview = "Invalid prompt. Please enter a description between 5 and 1000 characters. If this issue persists, please report it using the Feedback tab in the Help Guide."
                        success = False
                        logger.warning('{"event": "invalid_prompt_input", "session_id": "%s", "prompt_length": %d}', 
                                       session_id, len(prompt))
                
                if success:
                    formatted_html = format_document_html(preview)
                    response_placeholder.markdown(formatted_html, unsafe_allow_html=True)
                else:
                    response_placeholder.markdown(preview)
            except Exception as e:
                logger.error('{"event": "main_content_error", "session_id": "%s", "error": "%s"}', session_id, str(e), exc_info=True)
                response_placeholder.markdown("An unexpected error occurred. Please try again later. If this issue persists, please report it using the Feedback tab in the Help Guide.")
            finally:
                st.session_state.generation_in_progress = False

            st.session_state.messages.append({"role": "assistant", "content": preview})
            st.session_state.document_generated_successfully = success
            st.session_state.show_payment = success
            if success:
                st.session_state.current_document = preview
                st.session_state.show_download = True
                st.session_state.force_sidebar_open = True
                show_download_buttons()
                st.rerun()

def main():
    # --- Theme change detection for dynamic logo switching ---
    current_theme = st.get_option("theme.base")
    if "_last_theme" not in st.session_state:
        st.session_state["_last_theme"] = current_theme
    elif st.session_state["_last_theme"] != current_theme:
        st.session_state["_last_theme"] = current_theme
        st.experimental_rerun()
    # --- End theme change detection ---

    st.set_page_config(
        page_title="SmartClause",
        page_icon="assets/smartclause_badge.png",
        layout="centered"
    )
    
    init_session_state()

    try:
        is_dark_mode = st.get_option("theme.base") == "dark"
    except Exception:
        is_dark_mode = False

    if is_dark_mode:
        logo_path = "assets/smartclause_logo_light.png"
    else:
        logo_path = "assets/smartclause_logo_dark.png"

    st.logo(logo_path, size="medium")
    
    st.markdown("""
    <style>
     .title-container {
        text-align: center;
    }
    
    div[data-testid="stSidebar"] .payment-card {
        position: relative !important;
        bottom: auto;
        left: auto;
        background-color: rgb(240, 240, 240);
        padding: 20px;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        width: 100%;
        margin-bottom:  20px;
    }
    
    .stMarkdown, 
    .element-container div,
    .stChatMessage {
        -webkit-user-select: none !important;
        -moz-user-select: none !important;
        -ms-user-select: none !important;
        user-select: none !important;
    }
    
    .stMarkdown *, 
    .element-container div *,
    .stChatMessage * {
        -webkit-user-drag: none !important;
        -khtml-user-drag: none !important;
        -moz-user-drag: none !important;
        -o-user-drag: none !important;
        user-drag: none !important;
    }
    
    .stChatMessage div[data-testid="stMarkdownContainer"] {
        -webkit-user-select: none !important;
        -moz-user-select: none !important;
        -ms-user-select: none !important;
        user-select: none !important;
    }
    
    button[data-testid="StyledFullScreenButton"],
    .stMarkdown button {
        display: none !important;
    }
    
    .legal-document {
        background-color: white;
        border: 1px solid #ddd;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        padding: 30px !important;
        margin: 15px 0;
        border-radius: 8px;
    }
    
    .legal-document h1, .legal-document h2 {
        margin-bottom: 15px;
    }
    
    .legal-document p {
        text-align: justify;
    }
    
    .stChatMessage,
    .stChatMessageContent, 
    [data-testid="stChatMessageContent"],
    [data-testid="stMarkdownContainer"],
    .legal-document,
    .document-preview,
    .element-container,
    div[data-testid="stVerticalBlock"] > div {
        max-height: none !important;
        overflow: visible !important;
        height: auto !important;
    }

    .stMarkdown,
    .stChatMessage div[data-testid="stMarkdownContainer"] {
        max-height: none !important;
        overflow: visible !important;
    }

    * {
        max-height: none !important;
    }

    .legal-document {
        max-height: none !important;
        height: auto !important;
        overflow: visible !important;
        display: block;
        width: 100%;
    }

    .stImage > img {
        max-width: 300px;
        height: auto;
        margin: 0 auto;
        display: block;
    }

    @media (max-width: 768px) {
        .legal-document {
            padding: 20px;
            font-size: 14px;
        }
        .doc-paragraph {
            font-size: 14px;
        }
        .title-container {
            font-size: 24px;
        }
        .stImage > img {
            max-width: 100%;
        }
        .payment-card {
            padding: 15px;
        }
    }

    @media (max-width: 500px) {
        .legal-document {
            padding: 10px;
            font-size: 12px;
        }
        .doc-paragraph {
            font-size: 12px;
        }
        .stImage > img {
            max-width: 100%;
        }
    }
    </style>
    """, unsafe_allow_html=True)

    enhance_sidebar()
    show_main_content()

if __name__ == "__main__":
    main()